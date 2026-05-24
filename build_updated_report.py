from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:/College/03_Year_3/PhP/xampp/htdocs/BOOK_LOAN_MIDTERM")
OUT = Path(r"D:/College/03_Year_3/PhP/BaoCao_QuanLyThuVienSo_HCMUE_CapNhat.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "F2F4F7"
LIGHT = "F4F6F9"
MUTED = "555555"


def set_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def table_borders(table, color="B8C2CC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    spec = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }[level]
    for run in p.runs:
        set_run(run, size=spec[0], bold=True, color=spec[1])
    spacing(p, before=spec[2], after=spec[3])
    return p


def para(doc, text="", bold_prefix=None, align=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run(run, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            set_run(p.add_run(rest))
    else:
        set_run(p.add_run(text))
    if align is not None:
        p.alignment = align
    spacing(p)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        spacing(p, after=4, line=1.167)
        set_run(p.add_run(item))


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        spacing(p, after=4, line=1.167)
        set_run(p.add_run(item))


def add_table(doc, headers, rows, widths, caption=None, font_size=8.4):
    if caption:
        p = doc.add_paragraph()
        spacing(p, before=4, after=4)
        set_run(p.add_run(caption), size=10, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, GRAY)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=0)
        set_run(p.add_run(header), size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            spacing(p, after=0)
            if len(str(value)) <= 12 and idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=font_size)
    table_geometry(table, widths)
    para(doc)
    return table


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    table_borders(table, color="D6DEE8")
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    p = cell.paragraphs[0]
    spacing(p, after=3)
    set_run(p.add_run(title), bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    spacing(p, after=0)
    set_run(p.add_run(body), size=10)
    para(doc)


def add_image(doc, path, caption, width=5.9):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=2, after=8)
    set_run(p.add_run(caption), size=9, italic=True, color=MUTED)


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = sec.header.paragraphs[0]
    spacing(header, after=0)
    set_run(header.add_run("Hệ thống Quản lý Thư viện Số HCMUE - Báo cáo cập nhật"), size=9, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacing(footer, after=0)
    set_run(footer.add_run("Trang "), size=9, color=MUTED)
    page_field(footer)

    # Cover
    for text, size in (
        ("TRƯỜNG ĐẠI HỌC SƯ PHẠM THÀNH PHỐ HỒ CHÍ MINH", 14),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 13),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=4)
        set_run(p.add_run(text), size=size, bold=True, color=INK)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=20, after=8)
    set_run(p.add_run("BÁO CÁO ĐỒ ÁN CUỐI KÌ"), size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=6)
    set_run(p.add_run("MÔN HỌC LẬP TRÌNH PHP"), size=14, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=12, after=8)
    set_run(p.add_run("HỆ THỐNG QUẢN LÝ THƯ VIỆN SỐ HCMUE"), size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=18)
    set_run(p.add_run("Book Loan Midterm - Enterprise Digital Library & Co-working Ecosystem"), size=12, italic=True, color=MUTED)

    cover = [
        ("Nhóm thực hiện", "TTVP Group"),
        ("Thành viên", "Nguyễn Đỗ Thanh Phú - 49.01.104.111\nHuỳnh Hữu Tín - 49.01.104.153\nVũ Thành Vinh - 49.01.103.096\nNguyễn Trần Bảo Thái - 49.01.104.133"),
        ("Giảng viên hướng dẫn", "ThS. Võ Lê Phúc Hậu"),
        ("Phiên bản báo cáo", "Cập nhật theo mã nguồn và tài liệu dự án ngày 24/05/2026"),
    ]
    t = doc.add_table(rows=len(cover), cols=2)
    table_geometry(t, [2400, 5400], indent=800)
    table_borders(t)
    for row_idx, (key, value) in enumerate(cover):
        c0, c1 = t.rows[row_idx].cells
        shade(c0, GRAY)
        c0.text = ""
        c1.text = ""
        set_run(c0.paragraphs[0].add_run(key), bold=True, color=INK)
        for idx, line in enumerate(value.split("\n")):
            if idx:
                c1.paragraphs[0].add_run("\n")
            set_run(c1.paragraphs[0].add_run(line))
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Thành phố Hồ Chí Minh, tháng 5 năm 2026"), size=11, color=MUTED)
    doc.add_page_break()

    heading(doc, "LỜI MỞ ĐẦU", 1)
    para(doc, "Trong bối cảnh chuyển đổi số tại các trường đại học, thư viện không chỉ còn là nơi lưu trữ và cho mượn sách giấy, mà trở thành một hạ tầng học tập số kết hợp tài nguyên điện tử, không gian tự học, dữ liệu phân tích và các dịch vụ hỗ trợ cá nhân hóa. Hệ thống Quản lý Thư viện Số HCMUE được xây dựng nhằm mô phỏng một môi trường thư viện hiện đại cho sinh viên, thủ thư và quản trị viên.")
    para(doc, "So với phiên bản báo cáo ban đầu, hệ thống hiện tại đã được mở rộng đáng kể từ một ứng dụng quản lý mượn-trả cơ bản thành nền tảng thư viện số đầy đủ hơn: kiến trúc Backend API và Frontend SPA tách biệt, xác thực bằng Laravel Sanctum, phân quyền RBAC chi tiết, hàng đợi đặt chỗ, tài liệu số, phạt và thanh toán giả lập, đặt phòng tự học, thông báo, audit logs, báo cáo phân tích và tích hợp trợ lý AI Gemini.")
    para(doc, "Báo cáo này cập nhật lại toàn bộ nội dung học thuật và kỹ thuật cho phù hợp với hiện trạng mã nguồn Book Loan Midterm. Các phần mô tả chức năng, mô hình dữ liệu, API, kiến trúc triển khai, kiểm thử và hướng phát triển được viết lại theo cách bám sát repository hiện tại, đồng thời giữ cấu trúc báo cáo đồ án/học phần truyền thống.")

    heading(doc, "BẢNG PHÂN CÔNG THÀNH VIÊN NHÓM", 1)
    add_table(doc, ["MSSV", "Họ và tên", "Nội dung thực hiện cập nhật", "Mức độ hoàn thành", "Cộng tác"], [
        ["49.01.104.111", "Nguyễn Đỗ Thanh Phú", "Thiết kế giao diện React, triển khai Laravel API, tích hợp xác thực, báo cáo và tổng hợp tài liệu.", "100%", "A"],
        ["49.01.104.153", "Huỳnh Hữu Tín", "Đặc tả nghiệp vụ, kiểm thử luồng mượn-trả, phạt, thanh toán và quản lý thành viên.", "100%", "A"],
        ["49.01.103.096", "Vũ Thành Vinh", "Phân tích use case, hỗ trợ thiết kế tài liệu trình bày, rà soát chức năng AI và báo cáo.", "100%", "A"],
        ["49.01.104.133", "Nguyễn Trần Bảo Thái", "Thiết kế sơ đồ, đặc tả actor, room booking, tài liệu kiểm thử và slide thuyết trình.", "100%", "A"],
    ], [1500, 2100, 3900, 900, 840], font_size=8.0)

    heading(doc, "MỤC LỤC NỘI DUNG", 1)
    numbers(doc, [
        "Chương 1: Tổng quan đề tài",
        "Chương 2: Phân tích yêu cầu hệ thống",
        "Chương 3: Thiết kế hệ thống",
        "Chương 4: Triển khai hệ thống",
        "Chương 5: Kiểm thử và đánh giá",
        "Chương 6: Kết luận và hướng phát triển",
    ])
    doc.add_page_break()

    heading(doc, "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI", 1)
    heading(doc, "1.1. Bối cảnh và động lực", 2)
    para(doc, "Quy trình thư viện truyền thống thường phụ thuộc vào ghi nhận thủ công, xử lý tại quầy và các bảng dữ liệu rời rạc. Điều này gây khó khăn khi cần kiểm tra trạng thái sách theo thời gian thực, giới hạn số lượng mượn, nhắc hạn trả, xử lý phạt hoặc đánh giá nhu cầu sử dụng phòng học. Với nhóm người dùng là sinh viên đại học, yêu cầu truy cập tài nguyên từ xa, đọc tài liệu số, đặt chỗ và nhận gợi ý học tập ngày càng rõ rệt.")
    para(doc, "Book Loan Midterm hướng đến một hệ thống web hiện đại, trong đó frontend React đóng vai trò giao diện người dùng giàu tương tác, backend Laravel chịu trách nhiệm xác thực, phân quyền, nghiệp vụ và quản trị dữ liệu. Hệ thống được thiết kế theo tư duy API-first để dễ mở rộng sang mobile app, kiosk tại thư viện hoặc các dịch vụ tích hợp trong tương lai.")
    heading(doc, "1.2. Mục tiêu đề tài", 2)
    bullets(doc, [
        "Xây dựng hệ thống quản lý thư viện số cho HCMUE với đầy đủ vòng đời tra cứu, mượn, duyệt, nhận sách, trả sách và đặt chỗ.",
        "Hỗ trợ ba nhóm vai trò Student, Librarian và Admin với quyền hạn rõ ràng, tách biệt trách nhiệm vận hành và quản trị.",
        "Tích hợp tài liệu số, tiến trình đọc, đánh giá sách, danh sách yêu thích và đề xuất cá nhân hóa.",
        "Cung cấp module phạt và thanh toán giả lập để mô phỏng nghiệp vụ thực tế của thư viện.",
        "Quản lý phòng tự học/nhóm thông qua sơ đồ và quy tắc đặt phòng có thể cấu hình.",
        "Tăng tính chuyên nghiệp bằng queue worker, email nền, audit logs, OpenAPI documentation, Docker Compose và test suite.",
    ])
    heading(doc, "1.3. Phạm vi hệ thống", 2)
    para(doc, "Hệ thống hiện tại phục vụ ba nhóm người dùng chính. Sinh viên thao tác với danh mục sách, tài liệu số, yêu cầu mượn, phạt, thanh toán, phòng học, đánh giá và trợ lý AI. Thủ thư vận hành kho sách, xử lý mượn-trả, quản lý sinh viên, phòng học và giao dịch tiền phạt. Quản trị viên kế thừa quyền thủ thư và có thêm quyền cấu hình chính sách, miễn giảm phạt, quản lý nhân viên, báo cáo phân tích và xem nhật ký kiểm toán.")
    heading(doc, "1.4. Công nghệ sử dụng", 2)
    add_table(doc, ["Nhóm", "Công nghệ", "Vai trò trong hệ thống"], [
        ["Backend", "Laravel 12.x, PHP 8.2+", "REST API, controller, FormRequest, Resource, Eloquent ORM và middleware phân quyền."],
        ["Cơ sở dữ liệu", "SQLite local, MySQL 8.0 trong Docker", "Lưu domain data, trigger toàn vẹn, session, cache, queue và audit logs."],
        ["Xác thực", "Laravel Sanctum, Socialite OAuth", "Bearer token, đăng nhập email/mã định danh, Google/GitHub OAuth và quản lý phiên thiết bị."],
        ["Frontend", "React 19, TypeScript, Vite 6", "SPA cho student/staff/admin, route protection, AuthContext và domain API modules."],
        ["UI/UX", "TailwindCSS 4, Motion, lucide-react", "Giao diện responsive, dark/light mode, animation và icon nhất quán."],
        ["Định tuyến", "react-router-dom 7.x", "BrowserRouter, Routes, Navigate, Outlet; tương thích cách tổ chức route kiểu v6."],
        ["Đa ngôn ngữ", "react-i18next + Laravel SetLocale", "Chuyển đổi tiếng Việt/tiếng Anh ở cả giao diện và thông báo API."],
        ["QR/Thanh toán", "qrcode.react, @yudiel/react-qr-scanner", "Sinh mã QR, quét nhận sách, MoMo/VNPay sandbox và đối soát chuyển khoản."],
        ["AI", "Google Gemini API", "Chatbot, đề xuất cá nhân hóa, tạo tóm tắt và tag sách qua job nền."],
        ["Tài liệu API", "OpenAPI 3.0.3 + Swagger UI", "Cung cấp /api/openapi.json và /api/docs qua OpenApiController trong mã nguồn hiện tại."],
        ["DevOps", "Docker, Docker Compose", "Điều phối app, queue worker, MySQL db và frontend Nginx."],
        ["Kiểm thử", "PHPUnit, Vitest, Testing Library", "Feature tests backend và test giao diện/logic frontend."],
    ], [1400, 2600, 5360], font_size=8.0)
    callout(doc, "Ghi chú đối chiếu mã nguồn", "Một số mô tả ban đầu như Recharts hoặc Scramble được điều chỉnh trong báo cáo theo package và routes hiện tại: biểu đồ đang được triển khai bằng component React/CSS nội bộ, tài liệu API được phục vụ qua OpenApiController với chuẩn OpenAPI 3.0.3.")

    heading(doc, "CHƯƠNG 2: PHÂN TÍCH YÊU CẦU HỆ THỐNG", 1)
    heading(doc, "2.1. Tác nhân và phân quyền", 2)
    add_table(doc, ["Tác nhân", "Mô tả", "Quyền/tác vụ chính"], [
        ["Student", "Sinh viên sử dụng thư viện.", "Tìm kiếm sách, mượn/đặt chỗ, đọc tài liệu số, trả phí, đặt phòng, đánh giá, yêu thích, chat AI."],
        ["Librarian", "Nhân viên/thủ thư vận hành hằng ngày.", "Quản lý sách, sinh viên, mượn-trả, phòng học, thanh toán phạt và đối soát MoMo."],
        ["Admin", "Quản trị viên hệ thống.", "Có toàn bộ quyền thủ thư, quản lý thủ thư, cấu hình hệ thống, waive fines, xem audit logs và reports."],
        ["External Services", "Dịch vụ ngoài hệ thống.", "Gemini API, MoMo/VNPay sandbox, email server, OAuth providers và trình duyệt sinh viên."],
    ], [1300, 2200, 5860], font_size=8.0)
    para(doc, "RBAC được hiện thực bằng các bảng roles, permissions, role_permission, model_has_roles và model_has_permissions. Admin có 9 quyền cốt lõi; librarian mặc định có manage_books, manage_members, approve_requests và manage_rooms; student không có quyền quản trị trong bảng permissions mà được bảo vệ bằng role middleware.")
    heading(doc, "2.2. Danh sách chức năng hiện tại", 2)
    add_table(doc, ["Nhóm chức năng", "Student", "Librarian", "Admin"], [
        ["Tra cứu & catalog", "Tìm kiếm, lọc, autocomplete, xem review.", "CRUD sách, upload tài liệu số, import CSV.", "Giám sát và quản trị toàn bộ dữ liệu catalog."],
        ["Mượn/trả & đặt chỗ", "Gửi yêu cầu, hủy pending, xem lịch sử, vào hàng đợi khi hết sách.", "Duyệt/từ chối, QR pickup, trả sách, gia hạn, xử lý hàng đợi.", "Thiết lập giới hạn mượn và chính sách vận hành."],
        ["Digital Library", "Đọc PDF/web, download signed URL, lưu tiến trình đọc.", "Tải lên file số và quản lý metadata.", "Theo dõi tài nguyên số trong báo cáo."],
        ["Fines & Payments", "Xem phạt, MoMo QR/chuyển khoản, VNPay sandbox.", "Thu tiền mặt, duyệt/từ chối MoMo transfer, tạo phạt thủ công.", "Miễn giảm phạt và cấu hình mức phạt."],
        ["Study Rooms", "Đặt phòng, hủy, check-out.", "Quản lý phòng, duyệt/từ chối, check-in bằng mã.", "Cấu hình giới hạn phòng và xem thống kê sử dụng."],
        ["AI", "Chatbot và gợi ý sách cá nhân hóa.", "AI Enhancer tạo summary/tag khi thêm sách.", "Quản trị API key và theo dõi hiệu quả dữ liệu."],
        ["Bảo mật", "OTP, đổi mật khẩu, thiết bị đăng nhập, logout từ xa.", "Truy cập theo quyền vận hành.", "Audit logs, quản lý nhân sự và phân quyền."],
        ["Báo cáo", "Xem thống kê cá nhân cơ bản.", "Theo dõi vận hành hằng ngày.", "Dashboard, charts, CSV export, doanh thu phạt, tỷ lệ phòng."],
    ], [1700, 2500, 2500, 2660], font_size=7.4)

    heading(doc, "2.3. Yêu cầu chức năng theo module", 2)
    for subhead, items in [
        ("2.3.1. Module xác thực và tài khoản", [
            "Đăng nhập bằng email hoặc mã định danh, phân giải đúng bảng members hoặc librarians.",
            "Đăng ký sinh viên với OTP xác thực email; hỗ trợ quên mật khẩu và reset bằng OTP.",
            "OAuth qua provider bên ngoài, ghi nhận login_histories, danh sách thiết bị và revoke token từ xa.",
            "Rate limiting throttle:auth chống brute-force cho login/OTP/reset-password.",
        ]),
        ("2.3.2. Module sách, tài liệu số và tương tác cộng đồng", [
            "Danh mục sách vật lý và tài liệu số; hỗ trợ cover, vị trí, thể loại, số lượng tổng/còn, file format và lượt tải.",
            "Autocomplete và lọc theo search/category/resource purpose; digital files được truy cập qua signed URL.",
            "Favorites, rating/review sau khi hoàn thành mượn, reading_progress để đồng bộ trang đang đọc.",
            "AI metadata: tạo ai_summary và ai_tags phục vụ mô tả sách và gợi ý cá nhân hóa.",
        ]),
        ("2.3.3. Module mượn-trả, đặt chỗ và phạt", [
            "Vòng đời mượn gồm pending, approved, borrowed, returned, rejected và cancelled.",
            "Giới hạn số lượt mượn chủ động theo max_active_loans; không cho mượn trùng sách đang active.",
            "Khi sách hết, sinh viên có thể vào reservations queue; khi trả sách hệ thống tự chuyển cho người tiếp theo.",
            "Phạt quá hạn, hư hỏng, mất sách và phạt thủ công; fine có thể paid, unpaid hoặc waived.",
            "Thanh toán hỗ trợ tiền mặt, MoMo mock QR/chuyển khoản đối soát và VNPay sandbox.",
        ]),
        ("2.3.4. Module phòng tự học", [
            "Sinh viên đặt phòng theo ngày/giờ, số người và ghi chú; hệ thống chống trùng lịch và giới hạn theo policy.",
            "Thủ thư quản lý phòng, duyệt/từ chối booking, check-in bằng mã tại quầy và check-out.",
            "Tác vụ nền hủy no-show và hoàn tất booking đã kết thúc, giúp trạng thái phòng luôn cập nhật.",
        ]),
        ("2.3.5. Module quản trị, báo cáo và audit", [
            "Admin quản lý librarians, role/permission, settings, audit logs và CSV export reports.",
            "Reports tổng hợp xu hướng mượn/trả, doanh thu phạt, top sách và tỷ lệ sử dụng phòng.",
            "AuditLoggerService lưu action, target, old_values, new_values và thông tin request để truy vết thay đổi.",
        ]),
    ]:
        heading(doc, subhead, 3)
        bullets(doc, items)

    heading(doc, "2.4. Use case tiêu biểu", 2)
    add_table(doc, ["Mã", "Use case", "Tác nhân", "Kết quả chính"], [
        ["UC01", "Đăng nhập và quản lý phiên", "Student/Librarian/Admin", "Tạo Sanctum token, ghi login history, hỗ trợ revoke thiết bị."],
        ["UC02", "Tìm kiếm và xem chi tiết sách", "Student/Public", "Danh sách sách, autocomplete, review, trạng thái sẵn có."],
        ["UC03", "Gửi yêu cầu mượn sách", "Student", "Tạo borrowing pending hoặc reservation nếu hết sách."],
        ["UC04", "Duyệt và xác nhận nhận sách", "Librarian/Admin", "approved rồi borrowed sau QR pickup, cập nhật tồn kho."],
        ["UC05", "Trả sách và tính phạt", "Librarian/Admin", "returned, tăng tồn kho, sinh fine nếu quá hạn/hỏng/mất."],
        ["UC06", "Đặt phòng tự học", "Student", "Tạo room_booking theo policy và chống trùng lịch."],
        ["UC07", "Thanh toán phạt", "Student/Librarian", "MoMo/VNPay sandbox, tiền mặt, đối soát transfer."],
        ["UC08", "AI chatbot và gợi ý sách", "Student", "Trả lời tự nhiên, đề xuất liên kết sách phù hợp."],
        ["UC09", "AI Enhancer metadata", "Librarian/Admin", "Sinh summary/tag qua Gemini và job nền."],
        ["UC10", "Quản lý nhân sự và phân quyền", "Admin", "CRUD librarians, gán roles/permissions."],
        ["UC11", "Báo cáo phân tích", "Admin", "Biểu đồ, thống kê, CSV export."],
        ["UC12", "Audit logs", "Admin", "Tra cứu lịch sử hành động và diff JSON."],
    ], [760, 2300, 1900, 4400], font_size=7.8)
    heading(doc, "2.5. Yêu cầu phi chức năng", 2)
    add_table(doc, ["Tiêu chí", "Yêu cầu"], [
        ["Bảo mật", "Mật khẩu/OTP không lưu plaintext; protected endpoints dùng auth:sanctum; phân quyền role và permission ở backend."],
        ["Toàn vẹn dữ liệu", "DB transactions, lockForUpdate, SQLite triggers và foreign keys bảo vệ tồn kho, trạng thái mượn, booking."],
        ["Hiệu năng", "API phân trang/lọc; queue xử lý email/AI/report; cache database cho dữ liệu hot."],
        ["Khả dụng", "Docker Compose tách app, queue, db, frontend; health endpoint phục vụ giám sát."],
        ["Khả năng mở rộng", "REST API tách FE/BE, domain API modules, Eloquent models và services rõ trách nhiệm."],
        ["Khả năng kiểm thử", "PHPUnit feature tests cho nghiệp vụ backend; Vitest/Testing Library cho UI và logic frontend."],
        ["Quốc tế hóa", "Giao diện song ngữ Việt/Anh và API message theo Accept-Language."],
    ], [1800, 7560], font_size=8.2)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 3: THIẾT KẾ HỆ THỐNG", 1)
    heading(doc, "3.1. Kiến trúc tổng thể", 2)
    para(doc, "Hệ thống sử dụng kiến trúc tách biệt Backend API và Frontend SPA. Frontend chỉ giữ trạng thái trình bày, token và các dữ liệu tạm thời; mọi quyết định bảo mật, tồn kho, tiền phạt, đặt phòng và phân quyền đều được xác thực lại ở backend. Điều này tránh tình trạng tin vào client và giúp hệ thống dễ thay đổi giao diện hoặc mở rộng thêm client mới.")
    add_table(doc, ["Lớp", "Thành phần", "Trách nhiệm"], [
        ["Presentation", "React SPA, layouts, pages, AuthContext", "Hiển thị giao diện, route protection, i18n, gọi API qua src/api/client.ts."],
        ["API/Application", "Laravel controllers, FormRequests, Resources", "Validate request, điều phối service, trả JSON nhất quán."],
        ["Domain", "Eloquent models, services, middleware", "Nghiệp vụ mượn-trả, phạt, đặt phòng, RBAC, AI metadata."],
        ["Persistence", "SQLite/MySQL, migrations, triggers", "Lưu dữ liệu quan hệ, kiểm soát invariant và transaction."],
        ["Async", "Laravel Queue + Artisan commands", "Email OTP/nhắc hạn, cleanup booking, fine accrual, AI metadata."],
        ["External", "Gemini, OAuth, MoMo/VNPay sandbox", "AI, đăng nhập ngoài và mô phỏng thanh toán."],
    ], [1500, 2800, 5060], font_size=8.0)
    heading(doc, "3.2. Thiết kế cơ sở dữ liệu", 2)
    para(doc, "Cơ sở dữ liệu đã mở rộng từ 5 bảng lõi ban đầu sang một mô hình domain đầy đủ hơn, bao gồm người dùng, kho sách, mượn-trả, hàng đợi, đánh giá, tài liệu số, phạt, thanh toán, phòng học, thông báo, lịch sử đăng nhập, audit logs và RBAC.")
    add_table(doc, ["Bảng/Nhóm bảng", "Nội dung lưu trữ", "Ghi chú thiết kế"], [
        ["members", "Tài khoản sinh viên, OTP/email verified, OAuth, notification prefs.", "Đăng nhập qua Sanctum; có role student."],
        ["librarians", "Tài khoản thủ thư/admin.", "Gán roles/permissions polymorphic."],
        ["roles, permissions, role_permission, model_has_*", "RBAC chi tiết.", "Admin 9 quyền; librarian 4 quyền vận hành."],
        ["books", "Metadata sách, số lượng, tài liệu số, AI summary/tags.", "Soft delete; trigger bảo vệ quantity."],
        ["borrowing", "Vòng đời mượn sách.", "Trạng thái pending/approved/borrowed/returned/rejected/cancelled."],
        ["reservations", "Hàng đợi đặt chỗ khi hết sách.", "position tự cập nhật khi hủy/trả sách."],
        ["favorites, reviews, reading_progress", "Tương tác và tiến trình đọc.", "Dữ liệu đầu vào cho đề xuất AI."],
        ["library_settings", "Chính sách mượn, phạt, đặt phòng.", "Singleton id=1 với default constants trong model."],
        ["fines, fine_payments", "Phạt quá hạn/hư hỏng/mất sách và thanh toán.", "Hỗ trợ cash, MoMo, VNPay, waive và manual fines."],
        ["rooms, room_bookings", "Phòng tự học và lịch đặt.", "Trigger trạng thái booking, conflict checks."],
        ["notifications, login_histories, audit_logs", "Thông báo, thiết bị đăng nhập, nhật ký hệ thống.", "Tăng bảo mật và truy vết vận hành."],
        ["jobs, failed_jobs, job_batches, cache", "Hạ tầng queue/cache.", "Database driver cho môi trường local/Docker."],
    ], [2200, 3600, 3560], font_size=7.6)
    heading(doc, "3.3. Ràng buộc toàn vẹn và transaction", 2)
    bullets(doc, [
        "books_quantity_guard_insert và books_quantity_guard_update ngăn available_quantity âm hoặc vượt total_quantity.",
        "borrowing_guard_insert/update chỉ cho phép trạng thái hợp lệ và bảo vệ các trường bắt buộc.",
        "room_bookings_guard_insert/update bảo vệ trạng thái booking hợp lệ.",
        "Các thao tác duyệt mượn, xác nhận nhận sách, trả sách, hủy yêu cầu, hàng đợi reservation và booking phòng chạy trong DB::transaction().",
        "lockForUpdate() được dùng tại các điểm tranh chấp tài nguyên như tồn kho sách, fine, borrowing và room booking để giảm race condition.",
    ])
    heading(doc, "3.4. Thiết kế RBAC", 2)
    add_table(doc, ["Quyền", "Librarian", "Admin", "Mục đích"], [
        ["manage_books", "Có", "Có", "CRUD sách, import sách, upload tài liệu số."],
        ["manage_members", "Có", "Có", "CRUD/import tài khoản sinh viên."],
        ["approve_requests", "Có", "Có", "Duyệt, từ chối, QR pickup, trả sách, phạt và cash payment."],
        ["manage_rooms", "Có", "Có", "Quản lý phòng và booking."],
        ["waive_fines", "Không mặc định", "Có", "Miễn giảm tiền phạt với lý do."],
        ["manage_settings", "Không mặc định", "Có", "Cấu hình chính sách thư viện."],
        ["view_reports", "Không mặc định", "Có", "Xem dashboard và export báo cáo."],
        ["view_audit_logs", "Không mặc định", "Có", "Xem nhật ký kiểm toán hệ thống."],
        ["manage_librarians", "Không mặc định", "Có", "Quản lý thủ thư và gán quyền."],
    ], [2200, 1200, 1200, 4760], font_size=7.9)
    heading(doc, "3.5. Thiết kế REST API", 2)
    add_table(doc, ["Nhóm API", "Endpoint tiêu biểu", "Bảo vệ"], [
        ["Auth/OTP/OAuth", "POST /login, /register, /verify-otp, /forgot-password, GET /auth/{provider}/callback", "Public + throttle:auth"],
        ["Catalog/Digital", "GET /books, /books/autocomplete, /digital-documents, /books/{id}/reviews", "Public hoặc signed URL"],
        ["Student Borrowing", "POST /requests, GET /requests/me, DELETE /requests/{id}/cancel", "auth:sanctum + role:student"],
        ["Favorites/Reviews/Progress", "GET/POST /favorites, POST /books/{id}/reviews, PUT /reading-progress/{book}", "Student"],
        ["Fines/Payments", "GET /fines/me, POST /fines/{id}/momo/pay, /vnpay/pay, confirm-transfer", "Student; IPN public"],
        ["Room Booking", "GET /rooms, POST /room-bookings, admin room-bookings actions", "Public listing; student/admin protected"],
        ["Staff Operations", "members, books, requests, admin/fines, momo-pending", "role:admin,librarian + permission middleware"],
        ["Admin-only", "librarians, library-settings, reports, audit-logs, waive fines", "permission-based admin capability"],
        ["AI/OpenAPI/Health", "POST /ai/chat, GET /ai/recommendations, /openapi.json, /docs, /health", "AI protected; docs/health public"],
    ], [1900, 5160, 2300], font_size=7.6)
    heading(doc, "3.6. Thiết kế các workflow nghiệp vụ chính", 2)
    add_table(doc, ["Workflow", "Luồng chính", "Điểm kiểm soát"], [
        ["Mượn sách", "Student gửi yêu cầu -> pending; staff approve -> approved; QR pickup -> borrowed; return -> returned.", "Giới hạn active loans, tồn kho, duplicate active request, unpaid fine block."],
        ["Reservation queue", "Nếu sách hết, student vào hàng đợi; khi sách trả, hệ thống chuyển người đầu queue thành request/hold phù hợp.", "lockForUpdate, position shift, notification."],
        ["Fine accrual", "Command/service tính số ngày quá hạn theo fine_per_day, grace_period_days và max_fine_per_loan.", "Fine không vượt cap; paid/waived không tính lại."],
        ["MoMo/VNPay", "Student tạo payment -> nhận QR/redirect mock -> IPN/simulate/confirm transfer -> staff approve nếu cần.", "Gateway response JSON, status pending/completed/failed, audit trail."],
        ["Room booking", "Student chọn phòng/giờ/số người -> conflict policy -> pending/approved -> check-in -> completed/check-out.", "Max hours, max bookings/day, advance days, group size, check-in window."],
        ["AI metadata", "Staff nhấn generate hoặc tạo sách -> dispatch job -> Gemini sinh summary/tag -> lưu vào books.", "Fallback khi thiếu API key; queue ai-metadata."],
    ], [1800, 4660, 2900], font_size=7.6)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG", 1)
    heading(doc, "4.1. Backend Laravel", 2)
    para(doc, "Backend đặt tại thư mục BE/ và tuân theo cách tổ chức quen thuộc của Laravel: Controllers tiếp nhận request, FormRequests tách validation, Resources chuẩn hóa JSON response, Models định nghĩa quan hệ Eloquent, Services gom nghiệp vụ dùng lại, Middleware bảo vệ vai trò/quyền và Console Commands xử lý nền.")
    add_table(doc, ["Thành phần", "File/thư mục tiêu biểu", "Nội dung triển khai"], [
        ["Controllers", "app/Http/Controllers", "Auth, Book, Borrow, Fine, Room, Report, AI, Audit, Notification, Payment."],
        ["Requests", "app/Http/Requests", "Validation cho login, book upsert, room booking, settings, reset password, import."],
        ["Resources", "app/Http/Resources", "Ẩn trường nhạy cảm, format payload book/borrowing/member/user."],
        ["Models", "app/Models", "Member, Librarian, Book, Borrowing, Fine, RoomBooking, Role, Permission."],
        ["Services", "app/Services", "FineCalculationService, AuditLoggerService, BookAiMetadataService, BookCacheService."],
        ["Jobs/Mail/Notifications", "app/Jobs, app/Mail, app/Notifications", "GenerateBookAiMetadataJob, NotifyNewBookJob, OTP mail, due-soon mail, in-app notifications."],
        ["Commands", "app/Console/Commands", "CalculateOverdueFines, CleanupApprovedBorrowings, CleanupNoShowRoomBookings, SendDueSoonWarnings."],
    ], [1800, 2500, 5060], font_size=7.8)
    heading(doc, "4.2. Frontend React SPA", 2)
    para(doc, "Frontend đặt tại FE/book_loan/ và được xây dựng bằng React 19, TypeScript và Vite. Ứng dụng dùng AuthContext để phục hồi session từ localStorage, gắn Bearer token qua axios client, phân tách layout student/admin, hỗ trợ giao diện tối/sáng và song ngữ.")
    add_table(doc, ["Nhóm giao diện", "Màn hình tiêu biểu", "Chức năng nổi bật"], [
        ["Public/Auth", "Landing, Login, VerifyOtp, ForgotPassword, OAuthCallback", "Đăng nhập/đăng ký, OTP, OAuth, điều hướng theo role."],
        ["Student", "Home, Catalog, Digital, MyBooks, Requests, Favorites, Fines, RoomBooking, Settings", "Tra cứu, mượn, đọc tài liệu số, tiến trình đọc, thanh toán phạt, đặt phòng, quản lý phiên."],
        ["Admin/Librarian", "Dashboard, Inventory, Requests, Members, Fines, MomoApprovals, RoomBookings, Settings, Reports, AuditLogs, Librarians", "Vận hành thư viện, phê duyệt, báo cáo, cấu hình, nhân sự và audit."],
        ["Shared Components", "Header, Sidebar, AiChatbot, NotificationDropdown, ReadingRoom, Pagination, ErrorBoundary", "Trải nghiệm thống nhất, AI nổi, thông báo, xử lý lỗi và skeleton/empty states."],
        ["API Modules", "src/api/*.ts", "Tách theo domain: auth, books, borrow, fines, roomBooking, report, ai, audit."],
    ], [1800, 3300, 4260], font_size=7.6)
    # Screenshots are intentionally omitted in the final DOCX because Word COM
    # export can hang on this workstation when raster images are embedded.
    heading(doc, "4.3. Hệ thống thông báo và tác vụ nền", 2)
    bullets(doc, [
        "VerifyEmailOTP và ForgotPasswordOTP gửi mã xác thực trong các luồng tài khoản.",
        "DueSoonNotification và SendDueSoonWarnings nhắc sách sắp đến hạn, giảm rủi ro quá hạn.",
        "NewBookNotification và NotifyNewBookJob gửi thông báo khi nhập sách mới cho người bật nhận tin.",
        "GenerateBookAiMetadataJob chạy queue ai-metadata để không chặn phản hồi API khi tạo summary/tag.",
        "CleanupApprovedBorrowings hủy lượt approved nếu quá hạn pickup_deadline_hours; CleanupNoShowRoomBookings hủy no-show phòng học.",
    ])
    heading(doc, "4.4. Phạt, thanh toán và chính sách vận hành", 2)
    add_table(doc, ["Cấu hình", "Giá trị mặc định", "Ý nghĩa"], [
        ["loan_period_days", "14", "Số ngày mượn tiêu chuẩn."],
        ["max_active_loans", "5", "Số lượt active tối đa gồm pending/approved/borrowed."],
        ["fine_per_day", "5.000 VND", "Mức phạt quá hạn mỗi ngày."],
        ["max_fine_per_loan", "200.000 VND", "Trần phạt cho một lượt mượn."],
        ["grace_period_days", "0", "Số ngày ân hạn trước khi tính phạt."],
        ["damaged_book_fee", "50.000 VND", "Phạt cố định khi sách hư hỏng."],
        ["lost_book_fee", "200.000 VND", "Phạt cố định khi mất sách."],
        ["pickup_deadline_hours", "48", "Thời hạn nhận sách sau khi approved."],
        ["room_max_hours_per_booking", "3", "Số giờ tối đa cho một lượt đặt phòng."],
        ["room_max_bookings_per_day", "2", "Số lượt đặt phòng tối đa mỗi ngày."],
        ["room_advance_booking_days", "7", "Số ngày được phép đặt trước."],
        ["room_min_group_size", "2", "Số thành viên tối thiểu cho booking nhóm."],
        ["room_checkin_window_minutes", "15", "Cửa sổ check-in trước/sau giờ hẹn."],
        ["room_cancel_deadline_hours", "2", "Thời hạn tối thiểu để sinh viên tự hủy."],
    ], [3100, 1800, 4460], font_size=7.8)
    heading(doc, "4.5. AI Gemini", 2)
    para(doc, "AI được tích hợp theo hai hướng: hỗ trợ người dùng cuối qua chatbot/recommendations và hỗ trợ thủ thư qua AI Enhancer khi nhập sách. Khi thiếu GEMINI_API_KEY, backend vẫn có fallback để demo không bị ngắt hoàn toàn; khi có key, request được gửi tới Gemini 1.5 Flash qua Google Generative Language API.")
    add_table(doc, ["Tính năng AI", "Dữ liệu đầu vào", "Đầu ra"], [
        ["Chatbot tìm sách", "Câu hỏi tự nhiên, catalog sách, ngữ cảnh hội thoại.", "Câu trả lời thân thiện, gợi ý sách và đường dẫn chi tiết."],
        ["Đề xuất cá nhân hóa", "Lịch sử mượn, favorites, ratings, thể loại quan tâm.", "Danh sách sách phù hợp kèm lý do đề xuất."],
        ["AI metadata", "Title, author, genre, mô tả/tập tin số.", "ai_summary và ai_tags lưu vào books."],
    ], [2200, 3900, 3260], font_size=8.0)
    heading(doc, "4.6. DevOps và tài liệu API", 2)
    bullets(doc, [
        "Docker Compose điều phối 4 service: app, queue, db và frontend.",
        "Backend container expose http://localhost:8000; frontend Nginx expose http://localhost:5173.",
        "QUEUE_CONNECTION=database giúp chạy worker độc lập với request web.",
        "Health endpoint /api/health phục vụ kiểm tra DB, queue/cache và trạng thái ứng dụng.",
        "OpenAPI JSON tại /api/openapi.json và Swagger UI tại /api/docs giúp kiểm thử API trực tiếp.",
    ])

    doc.add_page_break()
    heading(doc, "CHƯƠNG 5: KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    heading(doc, "5.1. Chiến lược kiểm thử", 2)
    para(doc, "Hệ thống có hai lớp kiểm thử chính. Backend dùng PHPUnit Feature Tests để kiểm tra API, security, database invariants và workflow nghiệp vụ. Frontend dùng Vitest và Testing Library để kiểm tra route protection, login, UI trạng thái, admin pages, i18n, favorites, fines, digital library và room/autosave behavior.")
    add_table(doc, ["Nhóm test backend", "Nội dung bao phủ"], [
        ["AuthSecurityTest", "Rate limiting, login theo email, tài khoản chưa verify, role mismatch, expired token, 401/403 chuẩn hóa."],
        ["BorrowWorkflowTest", "Duplicate request, giới hạn active loans, duyệt/trả cập nhật tồn kho, reject reason, overdue fields."],
        ["DatabaseIntegrityTest", "Trigger tồn kho và ràng buộc khóa ngoại."],
        ["DigitalDocumentAccessTest", "Signed URL, filter tài liệu số, upload file hợp lệ/không hợp lệ."],
        ["BookReservationTest", "Đặt chỗ khi hết sách, position queue, hủy queue, return kích hoạt người kế tiếp."],
        ["FinePaymentWorkflowTest", "Fine schema, accrual, summary, cash payment, waive, unpaid fine block, MoMo transfer."],
        ["LibraryUpgradePhase2/3", "Audit logs, overdue warnings, AI metadata, reading progress, health endpoint, OpenAPI docs, queue tables."],
        ["AiFeatureTest, BookReviewTest, FavoriteBooksTest, LocalizationTest", "AI fallback, review rules, favorites, Accept-Language."],
    ], [2600, 6760], font_size=7.6)
    add_table(doc, ["Nhóm test frontend", "Nội dung bao phủ"], [
        ["login / protected-route", "Luồng đăng nhập, phân quyền route và chuyển hướng theo role."],
        ["admin-* tests", "Inventory upload, members, requests, settings, fines và report interactions."],
        ["catalog / digital / my-books", "Tìm kiếm, phân loại catalog, tài liệu số, cảnh báo quá hạn."],
        ["favorites / fines / room autosave", "Danh sách yêu thích, banner phạt, trang phạt và lưu tiến trình đọc/phòng."],
        ["i18n / theme / empty-state", "Đa ngôn ngữ, dark/light mode và trạng thái rỗng."],
    ], [2600, 6760], font_size=8.0)
    heading(doc, "5.2. Lệnh kiểm thử và build", 2)
    add_table(doc, ["Mục tiêu", "Lệnh", "Kết quả mong đợi"], [
        ["Backend routes", "cd BE && php artisan route:list", "Danh sách route public/protected đúng middleware."],
        ["Backend tests", "cd BE && php artisan test", "Feature tests pass, không phá vỡ invariant nghiệp vụ."],
        ["Frontend type/lint", "cd FE/book_loan && npm run lint", "TypeScript strict checks không lỗi."],
        ["Frontend tests", "cd FE/book_loan && npm run test", "Vitest suites pass."],
        ["Frontend build", "cd FE/book_loan && npm run build", "Sinh production bundle Vite thành công."],
        ["Docker smoke test", "docker-compose up --build -d", "4 service chạy, /api/health trả healthy."],
    ], [1900, 3300, 4160], font_size=7.6)
    callout(doc, "Lưu ý nghiệm thu", "Báo cáo mô tả bộ kiểm thử và lệnh kiểm thử hiện có trong repository. Khi chấm/demo chính thức, nhóm nên chạy lại toàn bộ lệnh trên môi trường sạch, chụp log pass và đính kèm vào phụ lục hoặc slide.")
    heading(doc, "5.3. Đánh giá hệ thống", 2)
    add_table(doc, ["Tiêu chí", "Đánh giá"], [
        ["Độ đầy đủ nghiệp vụ", "Bao phủ từ tra cứu, mượn/trả, đặt chỗ, tài liệu số, phòng học đến phạt/thanh toán và báo cáo."],
        ["Độ an toàn dữ liệu", "Có transaction, pessimistic locking, trigger và middleware backend; không đặt niềm tin vào client."],
        ["Tính mở rộng", "Tách FE/BE, API domain rõ, Docker và queue tạo nền cho triển khai lớn hơn."],
        ["Trải nghiệm người dùng", "SPA hiện đại, song ngữ, dark/light mode, notification, AI chatbot và nhiều màn hình chuyên biệt."],
        ["Tính chuyên nghiệp", "Có audit logs, OpenAPI docs, health endpoint, payment sandbox và test suite."],
    ], [2200, 7160], font_size=8.0)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    heading(doc, "6.1. Kết quả đạt được", 2)
    para(doc, "Book Loan Midterm đã phát triển vượt phạm vi hệ thống quản lý thư viện cơ bản ban đầu. Hệ thống hiện có kiến trúc full-stack rõ ràng, xác thực token, RBAC chi tiết, quản trị kho sách, workflow mượn-trả nhiều trạng thái, hàng đợi đặt chỗ, tài liệu số, đọc và lưu tiến trình, quản lý phạt/thanh toán, phòng tự học, thông báo, audit logs, báo cáo phân tích, tích hợp Gemini AI và môi trường Docker.")
    bullets(doc, [
        "Sinh viên có thể tự phục vụ nhiều nhu cầu: tìm sách, mượn/đặt chỗ, đọc tài liệu số, thanh toán phạt, đặt phòng và nhận đề xuất.",
        "Thủ thư có công cụ vận hành đầy đủ: duyệt yêu cầu, quét QR nhận sách, xử lý trả/gia hạn, import dữ liệu, đối soát phạt và quản lý phòng.",
        "Admin có khả năng kiểm soát hệ thống: phân quyền, chính sách, báo cáo, audit logs và miễn giảm phạt.",
        "Nền tảng kỹ thuật có queue, health check, OpenAPI, test suite và Docker giúp demo/triển khai ổn định hơn.",
    ])
    heading(doc, "6.2. Hạn chế hiện tại", 2)
    bullets(doc, [
        "MoMo/VNPay đang ở mức sandbox/mock, chưa phải tích hợp thanh toán production với đối soát ngân hàng thực tế.",
        "Gemini phụ thuộc GEMINI_API_KEY và chất lượng prompt/context; cần giám sát chi phí và bảo mật dữ liệu khi triển khai thật.",
        "Tài liệu số đã hỗ trợ đọc/lưu tiến trình nhưng các tính năng annotation, watermark cá nhân hóa và DRM vẫn là hướng mở rộng.",
        "Biểu đồ hiện tại dùng component nội bộ; có thể chuẩn hóa bằng thư viện chart chuyên dụng nếu yêu cầu trực quan phức tạp hơn.",
        "Hệ thống chưa triển khai multi-branch library, mobile app native và realtime push/WebSocket đầy đủ.",
    ])
    heading(doc, "6.3. Hướng phát triển", 2)
    add_table(doc, ["Hướng phát triển", "Mục tiêu"], [
        ["Thanh toán production", "Kết nối MoMo/VNPay thật, chữ ký bảo mật, webhook retry, reconciliation report."],
        ["AI nâng cao", "Semantic search bằng embedding, recommendation scoring kết hợp rule-based và AI, cache kết quả."],
        ["Digital reader nâng cao", "Highlight, ghi chú, bookmark, watermark, download quota và bảo vệ bản quyền."],
        ["Realtime notifications", "SSE/WebSocket cho trạng thái mượn, phòng, payment và thông báo sách mới."],
        ["CI/CD & monitoring", "GitHub Actions, log tập trung, alert queue/health, deploy staging/production."],
        ["Quản lý đa chi nhánh", "Bổ sung branches, scope dữ liệu theo cơ sở, chuyển sách giữa kho và báo cáo theo chi nhánh."],
    ], [2500, 6860], font_size=8.0)

    heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    numbers(doc, [
        "Mã nguồn dự án Book Loan Midterm, thư mục BE/ và FE/book_loan/, cập nhật ngày 24/05/2026.",
        "README.md và PROJECT_DOCUMENTATION.md của dự án Book Loan Midterm.",
        "Laravel Documentation: https://laravel.com/docs",
        "Laravel Sanctum Documentation: https://laravel.com/docs/sanctum",
        "React Documentation: https://react.dev",
        "Vite Documentation: https://vite.dev",
        "Tailwind CSS Documentation: https://tailwindcss.com/docs",
        "Google Gemini API Documentation: https://ai.google.dev",
        "Docker Documentation: https://docs.docker.com",
        "Vitest Documentation: https://vitest.dev",
        "PHPUnit Documentation: https://phpunit.de",
    ])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_margins(cell)
                for p in cell.paragraphs:
                    if p.paragraph_format.space_after is None:
                        spacing(p, after=0)
                    for run in p.runs:
                        if not run.font.name:
                            set_run(run)

    doc.core_properties.title = "Báo cáo cập nhật Hệ thống Quản lý Thư viện Số HCMUE"
    doc.core_properties.author = "TTVP Group"
    doc.core_properties.subject = "Book Loan Midterm - Laravel React Digital Library"
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_doc()
